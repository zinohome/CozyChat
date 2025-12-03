#!/usr/bin/env python3
"""
将 docs/ppt/ 目录下的 Markdown 文件转换为 Word 文档

使用方法:
    python scripts/convert_md_to_word.py
"""

import os
import sys
from pathlib import Path

# 添加项目根目录到路径
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

def convert_with_pypandoc(md_file: Path, output_dir: Path) -> bool:
    """使用 pypandoc 转换（需要系统安装 pandoc）"""
    try:
        import pypandoc
        
        output_file = output_dir / f"{md_file.stem}.docx"
        print(f"转换: {md_file.name} -> {output_file.name}")
        
        pypandoc.convert_file(
            str(md_file),
            'docx',
            outputfile=str(output_file),
            extra_args=['--standalone', '--toc']
        )
        return True
    except ImportError:
        print("pypandoc 未安装，尝试使用 python-docx...")
        return False
    except Exception as e:
        print(f"pypandoc 转换失败: {e}")
        return False

def convert_with_docx(md_file: Path, output_dir: Path) -> bool:
    """使用 python-docx + markdown 手动转换"""
    try:
        import markdown
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re
        
        # 读取 markdown 文件
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # 转换为 HTML
        html = markdown.markdown(
            md_content,
            extensions=['extra', 'codehilite', 'tables', 'toc']
        )
        
        # 创建 Word 文档
        doc = Document()
        
        # 解析 markdown 并转换为 Word
        lines = md_content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].rstrip()
            
            # 标题
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                p = doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                p = doc.add_heading(line[5:], level=4)
            elif line.startswith('##### '):
                p = doc.add_heading(line[6:], level=5)
            elif line.startswith('###### '):
                p = doc.add_heading(line[7:], level=6)
            # 分隔线
            elif line.startswith('---') or line.startswith('***'):
                doc.add_paragraph('─' * 50)
            # 代码块
            elif line.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                if code_lines:
                    p = doc.add_paragraph('\n'.join(code_lines))
                    p.style = 'Intense Quote'
            # 列表项
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\.\s', line):
                doc.add_paragraph(re.sub(r'^\d+\.\s', '', line), style='List Number')
            # 引用
            elif line.startswith('> '):
                p = doc.add_paragraph(line[2:], style='Intense Quote')
            # 空行
            elif not line.strip():
                doc.add_paragraph()
            # 普通段落
            else:
                # 处理内联格式
                text = line
                # 移除 markdown 链接格式 [text](url) -> text
                text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
                # 移除粗体 **text** -> text (保留文本)
                text = re.sub(r'\*\*([^\*]+)\*\*', r'\1', text)
                # 移除斜体 *text* -> text
                text = re.sub(r'\*([^\*]+)\*', r'\1', text)
                # 移除代码 `code` -> code
                text = re.sub(r'`([^`]+)`', r'\1', text)
                
                if text.strip():
                    p = doc.add_paragraph(text)
            
            i += 1
        
        # 保存文档
        output_file = output_dir / f"{md_file.stem}.docx"
        doc.save(str(output_file))
        print(f"✓ 转换成功: {output_file.name}")
        return True
        
    except ImportError as e:
        print(f"缺少依赖库: {e}")
        print("请安装: pip install python-docx markdown")
        return False
    except Exception as e:
        print(f"转换失败: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """主函数"""
    # 确定路径
    ppt_dir = project_root / 'docs' / 'ppt'
    output_dir = ppt_dir / 'word'
    
    if not ppt_dir.exists():
        print(f"错误: 目录不存在 {ppt_dir}")
        return
    
    # 创建输出目录
    output_dir.mkdir(exist_ok=True)
    
    # 获取所有 markdown 文件
    md_files = sorted(ppt_dir.glob('*.md'))
    
    if not md_files:
        print(f"未找到 markdown 文件在 {ppt_dir}")
        return
    
    print(f"找到 {len(md_files)} 个 markdown 文件")
    print(f"输出目录: {output_dir}\n")
    
    success_count = 0
    failed_files = []
    
    for md_file in md_files:
        print(f"\n处理: {md_file.name}")
        
        # 先尝试使用 pypandoc
        if convert_with_pypandoc(md_file, output_dir):
            success_count += 1
        # 如果失败，使用 python-docx
        elif convert_with_docx(md_file, output_dir):
            success_count += 1
        else:
            failed_files.append(md_file.name)
    
    # 输出结果
    print(f"\n{'='*60}")
    print(f"转换完成!")
    print(f"成功: {success_count}/{len(md_files)}")
    if failed_files:
        print(f"失败: {len(failed_files)} 个文件")
        for f in failed_files:
            print(f"  - {f}")
    print(f"\n输出目录: {output_dir}")

if __name__ == '__main__':
    main()

